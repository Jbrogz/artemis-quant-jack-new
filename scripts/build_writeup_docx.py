"""
build_writeup_docx.py
---------------------
Generates the comprehensive project-findings write-up for the Artemis Momentum
Factor Book study as a Microsoft Word .docx:

    docs/report/Artemis_Momentum_Findings.docx

This is an ASSEMBLY / PRESENTATION task. The analysis is complete and verified.
The document embeds verified literal narrative and table values from the
committed source artifacts; it does not recompute statistics. The highest-risk
Stage-2 and Stage-4 literal tables are covered by tests that compare them back
to their source markdown docs:

  - docs/report/Artemis_Momentum_Report.md            (narrative, primary source)
  - docs/STAGE2_RESULTS.md                             (full 21-variant table)
  - docs/STAGE4_RESULTS.md                             (gross/net, IS/OOS, capacity, regime)
  - docs/AUDIT.md                                      (universe + survivorship)
  - docs/specs/2026-05-30-artemis-momentum-design.md   (methodology + disclosures)
  - docs/report/figures/fig{1..5}_*.png                (figures)
  - scripts/build_report_figures.py                    (figure takeaway captions)

Dependency:  uv add python-docx
Run:         uv run python scripts/build_writeup_docx.py

The script also re-opens and validates the generated file (expected headings,
21-row significance table, 5 embedded images) and prints document statistics.
"""

from __future__ import annotations

import pathlib
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
ROOT = pathlib.Path(__file__).resolve().parent.parent
FIG_DIR = ROOT / "docs" / "report" / "figures"
OUT_PATH = ROOT / "docs" / "report" / "Artemis_Momentum_Findings.docx"

# Usable text width on US-Letter with 1" margins ≈ 6.5 in.
PAGE_TEXT_WIDTH_IN = 6.5

# ---------------------------------------------------------------------------
# Verdict / palette
# ---------------------------------------------------------------------------
VERDICT_RED = RGBColor(0xB2, 0x22, 0x22)
ACCENT_NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F3A5F"   # navy table-header fill
SURVIVOR_FILL = "E6F4EA"  # light-green highlight for the L5d row

BONFERRONI = "0.05 / 7 = 0.00714"

# ===========================================================================
# Source-of-truth content (quoted verbatim from the committed artifacts)
# ===========================================================================

# --- Full Stage-2 significance table: all 21 variants (docs/STAGE2_RESULTS.md) ---
# Columns mirror the source markdown table exactly.
STAGE2_COLUMNS = [
    "variant", "n", "naive t", "HAC t", "HLZ", "ann ret", "Sharpe (SE)",
    "autocorr", "span α", "span α t", "HAC p", "boot p",
    "disagree", "holds sign", "power", "DSR", "survives",
]

# Selection family (skip = 1) — the deployment candidates (7 rows).
STAGE2_FAMILY = [
    ["momentum_L14d_S1d", "69", "1.902", "1.945", "not_significant", "0.5128", "0.799 (0.425)", "no", "0.02343", "1.605", "0.0259", "0.0354", "no", "yes", "powered", "0.308", "no"],
    ["momentum_L1d_S1d", "69", "-1.247", "-1.300", "not_significant", "-0.3078", "-0.524 (0.422)", "no", "-0.02494", "-1.450", "0.9032", "0.9102", "no", "no", "powered", "0.000", "no"],
    ["momentum_L28d_S1d", "69", "0.911", "0.955", "not_significant", "0.2791", "0.383 (0.421)", "no", "-0.00166", "-0.088", "0.1698", "0.1780", "no", "yes", "powered", "0.060", "no"],
    ["momentum_L3d_S1d", "69", "1.606", "1.533", "not_significant", "0.4098", "0.674 (0.424)", "no", "0.01286", "0.822", "0.0626", "0.0656", "no", "no", "powered", "0.198", "no"],
    ["momentum_L56d_S1d", "69", "0.935", "0.756", "not_significant", "0.2683", "0.393 (0.525)", "yes", "0.00798", "0.360", "0.2249", "0.2024", "no", "no", "powered", "0.062", "no"],
    ["momentum_L5d_S1d", "69", "2.315", "2.403", "suggestive", "0.6163", "0.972 (0.428)", "no", "0.04651", "2.211", "0.0081", "0.0058", "yes", "no", "powered", "0.460", "yes"],
    ["momentum_L7d_S1d", "69", "2.068", "1.887", "not_significant", "0.6209", "0.868 (0.426)", "no", "0.04047", "1.621", "0.0296", "0.0214", "no", "yes", "powered", "0.352", "no"],
]

# Diagnostics (skip {2,3}) — reported, never eligible for deployment (14 rows).
STAGE2_DIAGNOSTICS = [
    ["momentum_L14d_S2d", "69", "2.369", "2.298", "suggestive", "0.6492", "0.995 (0.428)", "no", "0.03692", "2.205", "0.0108", "0.0140", "no", "yes", "powered", "0.481", "no"],
    ["momentum_L14d_S3d", "69", "3.241", "3.949", "significant", "0.8988", "1.361 (0.436)", "no", "0.06076", "3.892", "0.0000", "0.0006", "no", "yes", "powered", "0.805", "no"],
    ["momentum_L1d_S2d", "69", "0.075", "0.069", "not_significant", "0.0288", "0.032 (0.420)", "no", "-0.03388", "-1.491", "0.4724", "0.4459", "no", "no", "powered", "0.010", "no"],
    ["momentum_L1d_S3d", "69", "3.338", "2.851", "suggestive", "1.0432", "1.402 (0.437)", "no", "0.07494", "2.662", "0.0022", "0.0054", "no", "yes", "powered", "0.870", "no"],
    ["momentum_L28d_S2d", "69", "0.872", "0.937", "not_significant", "0.2497", "0.366 (0.421)", "no", "0.00329", "0.182", "0.1744", "0.1758", "no", "yes", "powered", "0.060", "no"],
    ["momentum_L28d_S3d", "69", "1.416", "1.393", "not_significant", "0.3889", "0.595 (0.423)", "no", "0.01942", "1.072", "0.0818", "0.0804", "no", "yes", "powered", "0.142", "no"],
    ["momentum_L3d_S2d", "69", "2.766", "2.378", "suggestive", "1.0109", "1.161 (0.505)", "yes", "0.04401", "2.116", "0.0087", "0.0192", "no", "no", "powered", "0.653", "no"],
    ["momentum_L3d_S3d", "69", "4.663", "5.015", "significant", "1.0807", "1.958 (0.452)", "no", "0.08106", "4.743", "0.0000", "0.0002", "no", "yes", "powered", "0.985", "no"],
    ["momentum_L56d_S2d", "69", "1.217", "0.990", "not_significant", "0.3780", "0.511 (0.522)", "yes", "0.00931", "0.425", "0.1610", "0.1466", "no", "no", "powered", "0.099", "no"],
    ["momentum_L56d_S3d", "69", "1.075", "0.891", "not_significant", "0.3159", "0.451 (0.513)", "yes", "0.01094", "0.528", "0.1866", "0.1658", "no", "no", "powered", "0.081", "no"],
    ["momentum_L5d_S2d", "69", "2.508", "2.588", "suggestive", "0.9616", "1.053 (0.429)", "no", "0.04526", "2.060", "0.0048", "0.0108", "yes", "yes", "powered", "0.558", "no"],
    ["momentum_L5d_S3d", "69", "2.857", "2.660", "suggestive", "0.8167", "1.200 (0.432)", "no", "0.05060", "2.187", "0.0039", "0.0030", "no", "yes", "powered", "0.675", "no"],
    ["momentum_L7d_S2d", "69", "2.097", "2.327", "suggestive", "0.6428", "0.881 (0.427)", "no", "0.03529", "1.691", "0.0100", "0.0116", "no", "yes", "powered", "0.372", "no"],
    ["momentum_L7d_S3d", "69", "1.925", "2.102", "suggestive", "0.5245", "0.808 (0.426)", "no", "0.03546", "2.082", "0.0178", "0.0110", "no", "yes", "powered", "0.330", "no"],
]

# --- Stage-4 gross/net, IS/OOS (docs/STAGE4_RESULTS.md) ---
S4_COLUMNS = ["segment", "gross Sharpe", "net Sharpe", "net ann ret",
              "net ann vol", "Sortino", "max DD", "Calmar", "hit rate",
              "ann turnover"]
S4_PRIMARY = [
    ["in-sample", "0.789", "0.664", "0.2490", "0.3752", "0.756", "-0.4851", "0.513", "0.574", "24.31"],
    ["out-of-sample", "0.897", "0.756", "0.2469", "0.3266", "1.031", "-0.2516", "0.981", "0.567", "28.22"],
]
S4_COMPARATOR = [
    ["in-sample", "0.073", "-0.031", "-0.0134", "0.4314", "-0.030", "-0.6809", "-0.020", "0.559", "23.90"],
    ["out-of-sample", "-0.158", "-0.270", "-0.0818", "0.3031", "-0.283", "-0.5168", "-0.158", "0.433", "26.57"],
]

S4_2X_COSTS = [
    ["1x costs", "0.789", "0.664", "0.2490", "0.3752", "0.756", "-0.4851", "0.513", "0.574", "24.31"],
    ["2x costs", "0.789", "0.552", "0.2072", "0.3753", "0.610", "-0.5219", "0.397", "0.574", "24.32"],
]

S4_LOOKBACK_COLUMNS = ["variant", "net Sharpe", "net ann ret", "note"]
S4_LOOKBACK = [
    ["momentum_L5d_S1d (chosen)", "0.664", "0.2490", "chosen lookback = 5d"],
    ["momentum_L2d_S1d", "-0.328", "-0.1155", "lookback-50% (skip/quantile unchanged)"],
    ["momentum_L8d_S1d", "0.621", "0.2645", "lookback+50% (skip/quantile unchanged)"],
]

S4_REGIME_COLUMNS = ["regime", "n", "mean net return"]
S4_REGIME = [
    ["bull", "17", "-0.00624"],
    ["bear", "28", "0.02195"],
    ["chop", "22", "0.03453"],
]

S4_EXTRA_COLUMNS = ["spec", "segment", "total return", "avg win", "avg loss"]
S4_EXTRA = [
    ["Primary L5d_S1d", "in-sample", "1.7225", "0.08584", "-0.06745"],
    ["Primary L5d_S1d", "out-of-sample", "0.6266", "0.07509", "-0.05136"],
    ["Comparator L28d_S1d", "in-sample", "-0.4518", "0.06534", "-0.08526"],
    ["Comparator L28d_S1d", "out-of-sample", "-0.2687", "0.06283", "-0.05991"],
]

# --- Universe / survivorship (docs/AUDIT.md) ---
AUDIT_COLUMNS = ["Metric", "Value"]
AUDIT_ROWS = [
    ["Artemis /asset catalog entries (as-of-today)", "1,013"],
    ["Assets with any price data (panel)", "846"],
    ["Date range", "2018-01-01 → 2026-05-30"],
    ["Panel rows (date × symbol)", "2,598,912"],
    ["Total terminal collapses (>90% drawdown, carried)", "238 (28%)"],
    ["— delisted cohort", "15"],
    ["— zombie cohort (still printing near-zero)", "223"],
    ["Left-censored assets (first price = 2018-01-01)", "272"],
    ["Assets ever eligible", "303"],
    ["Assets eligible on latest date (2026-05-30)", "226"],
]

# --- Figure captions: takeaways copied VERBATIM from scripts/build_report_figures.py ---
FIGURES = [
    (
        "fig1_cumulative_pnl.png",
        "Figure 1 — Cumulative P&L: L5d gross & net vs L28d net.",
        "L5d net grows steadily in-sample but is a regime-sensitive path; "
        "L28d net decays to near-zero — the canonical 4-week lookback does not survive costs.",
    ),
    (
        "fig2_drawdown.png",
        "Figure 2 — Drawdown (L5d net, primary candidate).",
        "In-sample max drawdown of −48.5% dwarfs the OOS −25.2%, consistent with "
        "the sign-instability finding; drawdown is material relative to returns at both horizons.",
    ),
    (
        "fig3_rolling_sharpe.png",
        "Figure 3 — Rolling 6-month Sharpe, L5d net.",
        "Rolling Sharpe oscillates widely and dips sharply negative in multiple sub-windows, "
        "confirming that any positive in-sample/OOS Sharpe is driven by episodic regimes rather "
        "than a stable factor.",
    ),
    (
        "fig4_variant_sharpe_bar.png",
        "Figure 4 — Gross Sharpe by variant: selection family (skip=1) vs diagnostics (skip=2,3).",
        "No selection-family variant clears the Bonferroni threshold on Newey-West; "
        "the eye-catching high Sharpes (L3d/S3, L14d/S3) are in the diagnostic zone "
        "and were never eligible for deployment.",
    ),
    (
        "fig5_variant_correlation.png",
        "Figure 5 — Factor-return correlation: selection family (skip=1).",
        "Short lookbacks (L1d–L5d) cluster together with high mutual correlation, "
        "while longer lookbacks (L14d–L56d) form a separate cluster; "
        "the family is far from orthogonal, reducing independent-test count below 7.",
    ),
]


# ===========================================================================
# Styling / low-level helpers
# ===========================================================================

def _set_cell_background(cell, hex_color: str) -> None:
    """Apply a solid shading fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold=False, color=None, size=8.0,
                   align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Replace a cell's text with a single styled run."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _style_base_font(document: Document) -> None:
    """Set a clean default body font."""
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12


def add_table(document: Document, columns, rows, *, header_size=8.0,
              body_size=8.0, highlight_first_col=False,
              highlight_row_predicate=None, col_widths=None):
    """
    Build a real Word table with a styled header row.

    highlight_row_predicate: optional fn(row_list) -> bool to shade a body row.
    """
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    hdr = table.rows[0].cells
    for i, col in enumerate(columns):
        _set_cell_text(hdr[i], str(col), bold=True, color=WHITE,
                       size=header_size, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_background(hdr[i], HEADER_FILL)

    # Body rows
    for row in rows:
        cells = table.add_row().cells
        shade = bool(highlight_row_predicate and highlight_row_predicate(row))
        for j, val in enumerate(row):
            align = (WD_ALIGN_PARAGRAPH.LEFT if j == 0
                     else WD_ALIGN_PARAGRAPH.CENTER)
            bold = highlight_first_col and j == 0
            _set_cell_text(cells[j], str(val), bold=bold, size=body_size,
                           align=align)
            if shade:
                _set_cell_background(cells[j], SURVIVOR_FILL)

    if col_widths is not None:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

    return table


def add_image_fit(document: Document, fig_name: str, caption: str,
                  takeaway: str) -> None:
    """Embed a figure sized to fit the page width, with a caption + takeaway."""
    path = FIG_DIR / fig_name
    if not path.exists():
        raise FileNotFoundError(
            f"Missing figure: {path} (run scripts/build_report_figures.py)"
        )
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(path), width=Inches(PAGE_TEXT_WIDTH_IN))

    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.bold = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = ACCENT_NAVY

    take = document.add_paragraph()
    take.alignment = WD_ALIGN_PARAGRAPH.CENTER
    take_run = take.add_run("Takeaway: " + takeaway)
    take_run.italic = True
    take_run.font.size = Pt(9)
    take_run.font.color.rgb = GREY


def add_bullets(document: Document, items) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            lead, body = item
            r = p.add_run(lead)
            r.bold = True
            p.add_run(body)
        else:
            p.add_run(item)


def add_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY


# ===========================================================================
# Document assembly
# ===========================================================================

def build_document() -> Document:
    doc = Document()
    _style_base_font(doc)

    # --- Title page header block ---
    title = doc.add_heading("Artemis Momentum Factor Book", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Project-Findings Write-Up — Cross-Sectional "
                          "Momentum on an Artemis-Sourced Crypto Spot Universe")
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = ACCENT_NAVY

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        "Date: 2026-05-30  |  Status: final deliverable (Stage 5)  |  "
        "Methodology: Project 1 Factor Book Guide"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = GREY

    # --- Verdict banner (one line) ---
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner.paragraph_format.space_before = Pt(10)
    banner.paragraph_format.space_after = Pt(10)
    b_run = banner.add_run(
        "VERDICT: Momentum is a rigorous, regime-dependent statistical NULL "
        "on the Artemis spot universe — NO-DEPLOY."
    )
    b_run.bold = True
    b_run.font.size = Pt(13)
    b_run.font.color.rgb = VERDICT_RED
    # Shade the banner paragraph for emphasis.
    p_pr = banner._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FBE9E7")
    p_pr.append(shd)

    prov = doc.add_paragraph()
    prov.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prov_run = prov.add_run(
        "This write-up is assembled from committed source artifacts. Key "
        "Stage-2 and Stage-4 tables are embedded as verified literals and "
        "checked against docs/STAGE2_RESULTS.md and docs/STAGE4_RESULTS.md; "
        "narrative and universe figures are verified literals from the "
        "committed docs. No new statistics are computed here."
    )
    prov_run.italic = True
    prov_run.font.size = Pt(8.5)
    prov_run.font.color.rgb = GREY

    # ---------------------------------------------------------------- 1 ----
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "We tested cross-sectional momentum on an Artemis-sourced crypto spot "
        "universe and find no deployable factor. The recommendation is: do not "
        "deploy."
    )
    add_bullets(doc, [
        ("What we tested. ",
         "Twenty-one momentum variants total — a pre-registered selection "
         "family of 7 lookbacks (1, 3, 5, 7, 14, 28, 56 days) at the convention "
         "skip of 1 day, plus 14 skip-{2,3} diagnostics reported for "
         "transparency but never eligible for deployment. The full Stage-2 "
         "significance battery was run on the in-sample slice; a sealed "
         "out-of-sample window was opened exactly once in the cost-aware "
         "backtest."),
        ("What held up — nothing, at the family level. ",
         "No selection-family variant is significant on the Newey-West HAC test "
         f"against the Bonferroni threshold ({BONFERRONI}). The strongest, "
         "momentum_L5d_S1d, is HLZ-“suggestive” with HAC t = 2.403 and "
         "HAC p = 0.0081 — above the Bonferroni threshold. It clears "
         "Bonferroni only through the bootstrap-override rule "
         "(bootstrap p = 0.0058), and is then disqualified outright by the "
         "Stage-2.6 deployment gate because it flips sign across subsamples "
         "(holds_sign = no). Its deflated Sharpe is 0.460 and the grid PBO is "
         "0.114."),
        ("Strongest variant, net of cost. ",
         "Characterized for completeness, momentum_L5d_S1d posts an in-sample "
         "net Sharpe of 0.664 (gross 0.789) and an out-of-sample net Sharpe of "
         "0.756 (gross 0.897). The OOS Sharpe did not collapse — but this is "
         "a single-regime artifact on a window of ~30 overlapping observations "
         "spent exactly once (the 2024 crypto bull), not a stable edge. The "
         "academic 4-week comparator momentum_L28d_S1d works only gross of "
         "costs: gross Sharpe 0.073 turns to net −0.031 in-sample and net "
         "−0.270 out-of-sample."),
        ("Recommendation. ",
         "Do not deploy momentum on the Artemis spot universe. The suggestive "
         "~5-day signal is a research lead, not a strategy. Capacity (~$36M) is "
         "comfortable and is not the binding constraint; the disqualification "
         "rests on multiple-testing failure, subsample sign-instability, regime "
         "dependence, fragility to the lookback parameter, and the disclosed "
         "data limitations."),
    ])

    # ---------------------------------------------------------------- 2 ----
    doc.add_heading("2. Objective & Scope", level=1)
    doc.add_paragraph(
        "Determine, with full statistical rigor and net of realistic trading "
        "costs, whether a cross-sectional momentum factor on an Artemis-sourced "
        "crypto spot universe has a true positive expected return — as "
        "opposed to a sample artifact. The null hypothesis for every test is a "
        "true mean factor return of zero. If momentum is insignificant, that "
        "honest null is the deliverable."
    )
    add_bullets(doc, [
        ("In scope: ",
         "the momentum factor only — construction, the statistical "
         "significance battery, in-sample deployment choice plus volatility "
         "targeting, a cost-aware backtest, and this research write-up."),
        ("Out of scope: ",
         "size, short-term reversal, betting-against-beta, and carry as deployed "
         "factors; any multi-factor combination. A small-minus-big size control "
         "is built solely as a test-only regressor inside the Stage-2.4 spanning "
         "regression — it is never formed as a portfolio, never deployed."),
        ("Data source: ",
         "all data sourced from Artemis. The study follows the Project 1 Factor "
         "Book Guide methodology strictly, differing only where the data source "
         "forces it; every such deviation is disclosed (Section 10)."),
    ])

    # ---------------------------------------------------------------- 3 ----
    doc.add_heading("3. Data & Universe", level=1)
    doc.add_paragraph(
        "The universe is enumerated programmatically from the Artemis /asset "
        "catalog (1,013 entries; 846 with any price data), on a daily "
        "point-in-time grid spanning 2018-01-01 to 2026-05-30 (2,598,912 panel "
        "rows). Eligibility filters (≥90 days history, ≥50% observation "
        "density, $10M market-cap floor, trailing-30d median 24H volume above "
        "threshold, staleness check, stablecoin and wrapped-token exclusion) are "
        "rebuilt as-of each date."
    )
    doc.add_paragraph(
        "Dead and collapsed coins are kept: 238 of 846 assets (28%) exhibit a "
        "terminal >90% drawdown (15 formally delisted + 223 still-printing "
        "“zombies”), and every one of their crash returns is carried "
        "into the P&L (the LUNA/lunc −99.99% collapse included). Recycled "
        "tickers are split into distinct synthetic assets so a revived ticker is "
        "never spliced onto a dead one. Holding return is the simple spot price "
        "return — Artemis exposes no perpetual funding, so no funding term is "
        "modeled (disclosed). Residual survivorship — coins fully purged from "
        "Artemis before the query — is an unestimable lower bound; reported "
        "factor returns are therefore upper bounds relative to a hypothetically "
        "complete universe."
    )
    add_table(doc, AUDIT_COLUMNS, AUDIT_ROWS, header_size=10, body_size=10,
              col_widths=[4.4, 2.1])
    add_caption(doc, "Table 1 — Universe and survivorship figures "
                     "(docs/AUDIT.md).")

    # ---------------------------------------------------------------- 4 ----
    doc.add_heading("4. Methodology", level=1)

    doc.add_heading("4.1 Momentum construction", level=2)
    doc.add_paragraph(
        "Signal = trailing log-sum return over a fixed lookback, shifted by a "
        "skip. The 7-lookback grid is fixed in advance (academic 14/28/56-day "
        "plus crypto-short 1/3/5/7-day); skip is a nuisance parameter fixed by "
        "convention at 1 day. The eligible universe is sorted into quintiles, "
        "long the top 20% / short the bottom 20%, equal-weight within each leg, "
        "dollar-neutral. The signal uses data through the close of t; the "
        "position is entered at the close of t+1 (a mandatory, unit-tested "
        "one-period lag — no look-ahead)."
    )

    doc.add_heading("4.2 Stage-2 significance battery", level=2)
    doc.add_paragraph("Per variant, the battery runs:")
    add_bullets(doc, [
        ("Newey-West HAC t-stat ",
         "(autocorrelation-robust) is the reported mean-return test, with "
         "bandwidth covering the holding-period overlap; the naive t-stat is "
         "shown but flagged biased and is never the headline."),
        ("Lo (2002) Sharpe SE ",
         "(autocorrelation-corrected when Ljung-Box fires)."),
        ("Spanning regression ",
         "of the factor on {equal-weight market return, small-minus-big size "
         "control} reporting the HAC-t alpha (the size control is a TEST-ONLY "
         "regressor)."),
        ("Bonferroni & Harvey-Liu-Zhu ",
         f"on the pre-registered family of 7 ({BONFERRONI}) and HLZ tiers."),
        ("Stationary block bootstrap ",
         "(arch) as the bootstrap of record; a Newey-West/bootstrap "
         "disagreement across the adjusted threshold defers to the bootstrap."),
        ("Subsample sign-stability ",
         "(halves and thirds) as a hard deployment gate — a sign-flip "
         "disqualifies regardless of t-stat."),
        ("Deflated Sharpe (DSR) and PBO/CSCV ",
         "to quantify selection overfit."),
    ])
    doc.add_paragraph(
        "Sealed OOS and cost-aware backtest. The most recent ~30% of rebalance "
        "dates (OOS_START = 2023-12-02) were reserved before any statistic was "
        "computed and opened exactly once, in Stage 4, under a single-use guard. "
        "The backtest applies a 10 bps spot taker fee per side plus size-scaled "
        "tiered slippage, executes at the t+1 close, and reports the full net "
        "metric set plus a capacity estimate."
    )

    # ---------------------------------------------------------------- 5 ----
    doc.add_heading("5. Significance Results", level=1)
    doc.add_paragraph(
        "No selection-family (skip = 1) variant is significant on the "
        f"Newey-West HAC test at the Bonferroni threshold ({BONFERRONI}). The "
        "lowest family p-value is momentum_L5d_S1d at 0.0081 — above the "
        "0.00714 threshold. L5d is HLZ-“suggestive” but disqualified by "
        "sign-instability (holds_sign = no) and only clears Bonferroni via the "
        "bootstrap-override rule (bootstrap p = 0.0058). The visually striking "
        "t≈5 diagnostics (e.g. L3d/S3d HAC t = 5.015, L14d/S3d HAC t = 3.949) "
        "are NOT in the selection family and were never eligible for deployment."
    )
    doc.add_paragraph(
        "Was there alpha? Only momentum_L5d_S1d has a spanning-alpha t-stat above "
        "2 (2.211) within the selection family; all other family variants are "
        "below 1.63 (the next-highest is L7d at 1.621, then L14d at 1.605). There "
        "is no robust alpha. Grid-level overfit controls: PBO/CSCV = 0.114 (below "
        "the 0.5 threshold but on a correlated family — Figure 5 shows short "
        "lookbacks cluster tightly, so the effective independent-test count is "
        "below 7)."
    )

    doc.add_heading("5.1 Selection family (skip = 1) — deployment candidates",
                    level=2)
    add_table(
        doc, STAGE2_COLUMNS, STAGE2_FAMILY, header_size=6.0, body_size=6.0,
        highlight_first_col=True,
        highlight_row_predicate=lambda r: r[0] == "momentum_L5d_S1d",
    )
    add_caption(doc, "Table 2 — Selection-family significance battery "
                     "(docs/STAGE2_RESULTS.md). The shaded row is the strongest "
                     "candidate, momentum_L5d_S1d: the sole bootstrap-override "
                     "Bonferroni survivor, disqualified by the §2.6 "
                     "sign-stability gate (holds sign = no).")

    doc.add_heading("5.2 Diagnostics (skip = 2, 3) — reported, not selected",
                    level=2)
    doc.add_paragraph(
        "These 14 variants are reported for transparency and were never eligible "
        "for deployment. Together with the 7-variant family they make up the full "
        "21-variant trial grid."
    )
    add_table(doc, STAGE2_COLUMNS, STAGE2_DIAGNOSTICS, header_size=6.0,
              body_size=6.0, highlight_first_col=True)
    add_caption(doc, "Table 3 — Diagnostic variants (docs/STAGE2_RESULTS.md). "
                     "The high diagnostic Sharpes (L3d/S3d, L14d/S3d) sit outside "
                     "the pre-registered selection family.")

    # ---------------------------------------------------------------- 6 ----
    doc.add_heading("6. Backtest & Out-of-Sample", level=1)
    doc.add_paragraph(
        "For completeness we characterized the strongest candidate, "
        "momentum_L5d_S1d, net of realistic spot costs. This characterizes the "
        "candidate; it does not rescue it — the Stage-2 disqualification "
        "stands."
    )

    doc.add_heading("6.1 Gross vs net, in-sample vs out-of-sample", level=2)
    doc.add_paragraph("Primary candidate momentum_L5d_S1d:")
    add_table(doc, S4_COLUMNS, S4_PRIMARY, header_size=8.0, body_size=8.0,
              highlight_first_col=True)
    add_caption(doc, "Table 4 — Primary L5d_S1d gross/net, IS/OOS "
                     "(docs/STAGE4_RESULTS.md).")
    doc.add_paragraph(
        "Academic 4-week comparator momentum_L28d_S1d — “works only "
        "gross”: positive gross in-sample (Sharpe 0.073) but net-negative "
        "both in-sample (net −0.031) and out-of-sample (net −0.270, "
        "gross −0.158). The canonical horizon does not work net of costs."
    )
    add_table(doc, S4_COLUMNS, S4_COMPARATOR, header_size=8.0, body_size=8.0,
              highlight_first_col=True)
    add_caption(doc, "Table 5 — Comparator L28d_S1d gross/net, IS/OOS "
                     "(docs/STAGE4_RESULTS.md).")
    add_bullets(doc, [
        ("IS − OOS net Sharpe gap = −0.092. ",
         "The OOS figure exceeds in-sample, but this is a spent-once, "
         "~30-observation, single-regime (2024-bull) artifact — not "
         "evidence of a stable edge. A near-zero or negative OOS Sharpe would "
         "have been reported as overfitting; the positive figure is reported "
         "as-is and interpreted as regime exposure."),
        ("Capacity ≈ $36,044,481 ",
         "— the AUM at which size-scaled slippage on the actual "
         "per-rebalance traded order (~2.0× summed one-way turnover) erases "
         "the gross edge (per-rebalance gross edge 0.02436). This is comfortably "
         "above a $1M book, so capacity is not the binding constraint."),
        ("Turnover ",
         "is ~24 (IS) / ~28 (OOS) annualized — high, consistent with a "
         "5-day signal."),
    ])
    doc.add_paragraph(
        "Additional net metrics (total return, average win / loss) are reported "
        "in the Appendix (Table 9)."
    )

    doc.add_heading("6.2 Robustness", level=2)
    doc.add_paragraph(
        "2× costs (in-sample, primary): net Sharpe degrades from 0.664 to "
        "0.552 (gross unchanged at 0.789; max DD widens to −0.5219). The "
        "edge thins materially under a doubled cost assumption. This is a "
        "sensitivity rerun of the deployed construction, not a re-selection."
    )
    add_table(doc, S4_COLUMNS, S4_2X_COSTS, header_size=8.0, body_size=8.0,
              highlight_first_col=True)
    add_caption(doc, "Table 6 — 2×-cost sensitivity, in-sample "
                     "(docs/STAGE4_RESULTS.md).")
    doc.add_paragraph(
        "±50% lookback (in-sample, net): the construction is fragile to its "
        "one free parameter. A −50% lookback (L2d) flips the candidate "
        "net-negative (−0.328) — not robust."
    )
    add_table(doc, S4_LOOKBACK_COLUMNS, S4_LOOKBACK, header_size=9.0,
              body_size=9.0, highlight_first_col=True,
              col_widths=[2.7, 1.2, 1.2, 1.4])
    add_caption(doc, "Table 7 — ±50% lookback robustness, in-sample net "
                     "(docs/STAGE4_RESULTS.md).")
    doc.add_paragraph(
        "Regime breakdown (in-sample, net mean return per regime):"
    )
    add_table(doc, S4_REGIME_COLUMNS, S4_REGIME, header_size=9.0, body_size=9.0,
              highlight_first_col=True, col_widths=[2.2, 1.0, 2.0])
    add_caption(doc, "Table 8 — Regime breakdown, in-sample net mean return "
                     "(docs/STAGE4_RESULTS.md).")
    cav = doc.add_paragraph()
    cav_run = cav.add_run(
        "Caveat (do not over-read): this is a full-sample, descriptive "
        "partition, not a walk-forward signal — it could not have been "
        "traded ex-ante. The chop bucket is the top-|market-move| tercile, which "
        "on this sample skews toward large up moves, so the apparent "
        "“negative in bull / positive in bear” contrast is overstated "
        "and is an artifact of where the magnitude cut falls. The disqualifying "
        "evidence is the Stage-2 §2.6 sign-instability itself, not this "
        "descriptive split."
    )
    cav_run.italic = True
    cav_run.font.size = Pt(9.5)
    cav_run.font.color.rgb = GREY
    doc.add_paragraph(
        "One-shot OOS, single-regime character. The OOS window is 30 "
        "overlapping-regime observations spent exactly once; a single favorable "
        "stretch (the 2024 crypto bull) carries the positive OOS Sharpe, "
        "interpreted as regime exposure, consistent with the sign-flip "
        "disqualification."
    )

    # ---------------------------------------------------------------- 7 ----
    doc.add_heading("7. Figures", level=1)
    doc.add_paragraph(
        "The five figures below visualize the P&L path, drawdown, rolling Sharpe, "
        "per-variant Sharpe, and family correlation. Each takeaway is copied "
        "verbatim from the figure-generation script."
    )
    for fig_name, caption, takeaway in FIGURES:
        add_image_fit(doc, fig_name, caption, takeaway)

    # ---------------------------------------------------------------- 8 ----
    doc.add_heading("8. Conclusion & Recommendation", level=1)
    doc.add_paragraph(
        "Do not deploy. Target allocation to a momentum sleeve on the Artemis "
        "spot universe: zero."
    )
    add_bullets(doc, [
        ("Expected Sharpe (range): ",
         "for the strongest candidate, net Sharpe sits in roughly 0.66–0.76 "
         "in the realized sample, but with no statistically reliable edge and a "
         "true expected Sharpe indistinguishable from zero once multiple testing, "
         "sign-instability, and the single-regime OOS are accounted for. We do "
         "not represent this as a deployable expectation."),
        ("The suggestive ~5-day signal is a research lead, not a strategy. ",
         "It would warrant follow-up only with (a) more independent observations, "
         "(b) a cleaner, lower-turnover construction, and (c) a true "
         "point-in-time universe that addresses residual survivorship."),
        ("Key risks if (against this recommendation) deployed: ",
         "regime dependence (the return is regime exposure, not a stable factor); "
         "selection / multiple-testing fragility (no family Bonferroni survivor "
         "on its own terms); subsample sign-instability (the hard disqualifier); "
         "parameter fragility (±50% lookback breaks it); and data "
         "limitations (no funding, daily t+1-close execution, volume "
         "unreliability, residual survivorship)."),
        ("Next steps: ",
         "treat momentum as closed for deployment on this universe; if revisited, "
         "pursue the above three conditions, and consider it only as one input to "
         "a broader, properly multiple-testing-controlled research program — "
         "never as a standalone sleeve."),
    ])

    # ---------------------------------------------------------------- 9 ----
    doc.add_heading("9. Limitations & Disclosures", level=1)
    add_bullets(doc, [
        ("No funding / spot: ",
         "Artemis exposes no perpetual funding; returns and costs carry no "
         "funding term (the guide's third cost component is N/A and stated)."),
        ("Daily, t+1-close execution: ",
         "Artemis serves daily granularity only; fills are at the t+1 close with "
         "slippage applied — a one-period lag, not a costless intraday "
         "fill."),
        ("Wrapped-token exclusion: ",
         "wrapped tokens (WBTC, WETH, stETH, …) are excluded as redundant "
         "price exposures of their underlyings — a deliberate, disclosed "
         "deviation from the guide."),
        ("Residual (as-of-today catalog) survivorship: ",
         "the Artemis catalog is an as-of-today snapshot with no "
         "listing/delisting dates; coins fully purged before the query are "
         "unrecoverable. This is an unestimable lower bound on survivorship bias; "
         "reported factor returns are upper bounds."),
        ("Volume reliability: ",
         "only 24H_VOLUME is historical (30D_VOLUME is real-time only) and is "
         "unreliable cross-sectionally; liquidity is gated on a market-cap floor "
         "+ trailing-30d median 24H volume — a disclosed deviation from the "
         "guide's mean-ADV rule."),
        ("Survivorship flows into P&L: ",
         "a collapsed short-leg coin's crash books as a positive contribution; "
         "dead coins are not dropped (238/846, 28% terminal collapses carried)."),
        ("Gross vs net, never flattered: ",
         "spot costs reduce every reported Sharpe; net is never flattered above "
         "gross. The OOS window was opened exactly once (single-use guard)."),
    ])

    # ---------------------------------------------------------------- 10 ---
    doc.add_heading("10. Appendix", level=1)

    doc.add_heading("10.1 Additional Stage-4 net metrics", level=2)
    add_table(doc, S4_EXTRA_COLUMNS, S4_EXTRA, header_size=9.0, body_size=9.0,
              highlight_first_col=True, col_widths=[1.9, 1.4, 1.2, 1.0, 1.0])
    add_caption(doc, "Table 9 — Additional §4.5 net metrics: total "
                     "return, average win / loss (docs/STAGE4_RESULTS.md).")

    doc.add_heading("10.2 Methodology references", level=2)
    add_bullets(doc, [
        "Newey, W.K. & West, K.D. (1987) — HAC (autocorrelation-robust) "
        "standard errors, the reported mean-return test.",
        "Lo, A.W. (2002) — the statistics of Sharpe ratios "
        "(autocorrelation-corrected Sharpe standard error).",
        "Harvey, Liu & Zhu (2016) — “...and the Cross-Section of "
        "Expected Returns” (multiple-testing thresholds / HLZ tiers).",
        "Bailey & Lopez de Prado — the Deflated Sharpe Ratio and PBO/CSCV "
        "(probability of backtest overfitting).",
        "Politis & Romano — the stationary block bootstrap (implemented via "
        "the arch library), the bootstrap of record.",
        "Project 1 Factor Book Guide — authoritative construction, "
        "significance, deployment, and backtest methodology.",
    ])

    doc.add_heading("10.3 How to reproduce", level=2)
    doc.add_paragraph(
        "The pipeline is deterministic and driven by scripts in scripts/ "
        "(uv-managed). The stages run in order:"
    )
    add_bullets(doc, [
        ("make probe ", "— live Artemis connectivity + coverage probe "
                        "(scripts/probe_artemis.py)."),
        ("make universe ", "— survivorship-corrected universe panel "
                          "(scripts/build_universe.py)."),
        ("uv run python scripts/build_returns.py ",
         "— holding returns (crash-carry)."),
        ("uv run python scripts/build_factor_returns.py ",
         "— the 21-variant momentum factor grid."),
        ("uv run python scripts/run_stage2.py ",
         "— the Stage-2 significance battery → data/stats/"
         "significance.parquet."),
        ("uv run python scripts/run_backtest.py ",
         "— the cost-aware, sealed-OOS backtest → data/backtest/."),
        ("uv run python scripts/build_report_figures.py ",
         "— the five report figures → docs/report/figures/."),
        ("uv run python scripts/build_report.py ",
         "— the deterministic research-report PDF."),
        ("uv run python scripts/build_writeup_docx.py ",
         "— this findings write-up (.docx)."),
    ])
    closing = doc.add_paragraph()
    closing_run = closing.add_run(
        "End of write-up. All figures and statistics are reproduced verbatim "
        "from the committed source-of-truth artifacts; the verdict is NO-DEPLOY."
    )
    closing_run.italic = True
    closing_run.font.size = Pt(9)
    closing_run.font.color.rgb = GREY

    return doc


# ===========================================================================
# Validation
# ===========================================================================

def validate(path: pathlib.Path) -> dict:
    """Re-open the .docx and assert structural expectations."""
    doc = Document(str(path))

    headings = [p.text for p in doc.paragraphs
                if p.style.name.startswith("Heading")]
    # Top-level sections are the Heading 1 paragraphs "1." .. "10." (not "4.1").
    h1 = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]

    expected_h1_starts = [f"{i}." for i in range(1, 11)]
    for start in expected_h1_starts:
        assert any(h.startswith(start) for h in h1), \
            f"Missing top-level section starting '{start}'. Found: {h1}"
    assert len(h1) == 10, f"Expected 10 top-level sections, found {len(h1)}."

    # The 21-row significance grid is the union of the family (7) + diagnostics
    # (14) tables; each has a header + body. Verify both tables exist with the
    # right body-row counts, and that the union equals 21 variant rows.
    n_cols = len(STAGE2_COLUMNS)
    family_tbl = None
    diag_tbl = None
    for t in doc.tables:
        if len(t.columns) == n_cols:
            body = len(t.rows) - 1  # minus header
            if body == 7 and family_tbl is None:
                family_tbl = t
            elif body == 14 and diag_tbl is None:
                diag_tbl = t
    assert family_tbl is not None, "Missing 7-row selection-family table."
    assert diag_tbl is not None, "Missing 14-row diagnostics table."
    total_variant_rows = (len(family_tbl.rows) - 1) + (len(diag_tbl.rows) - 1)
    assert total_variant_rows == 21, \
        f"Expected 21 variant rows total, found {total_variant_rows}."

    # 5 embedded images.
    image_parts = [p for p in doc.part.package.iter_parts()
                   if "media" in p.partname]
    n_images = len(image_parts)
    assert n_images == 5, f"Expected 5 embedded images, found {n_images}."

    # Verdict banner present.
    assert any("NO-DEPLOY" in p.text for p in doc.paragraphs), \
        "Verdict banner (NO-DEPLOY) not found."

    # Spot-check verbatim numbers in cells.
    family_cells = [c.text for row in family_tbl.rows for c in row.cells]
    for token in ["momentum_L5d_S1d", "2.403", "0.0081", "0.0058", "2.211",
                  "0.460"]:
        assert token in family_cells, f"Expected verbatim token '{token}' " \
            "missing from family table."

    return {
        "n_paragraphs": len(doc.paragraphs),
        "n_headings": len(headings),
        "n_top_level_sections": len(h1),
        "n_tables": len(doc.tables),
        "n_variant_rows": total_variant_rows,
        "n_images": n_images,
        "headings": headings,
    }


# ===========================================================================
# Main
# ===========================================================================

def main() -> int:
    print("=== Artemis Momentum — Findings Write-Up (.docx) ===")
    print(f"Output: {OUT_PATH}\n")

    doc = build_document()
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"  saved: {OUT_PATH}")

    print("\n--- Validation (re-open + assert) ---")
    stats = validate(OUT_PATH)
    print(f"  paragraphs:           {stats['n_paragraphs']}")
    print(f"  headings:             {stats['n_headings']}")
    print(f"  top-level sections:   {stats['n_top_level_sections']}")
    print(f"  tables:               {stats['n_tables']}")
    print(f"  significance rows:    {stats['n_variant_rows']} (7 family + 14 diagnostics)")
    print(f"  embedded images:      {stats['n_images']}")
    size_kb = OUT_PATH.stat().st_size / 1024
    print(f"  file size:            {size_kb:.1f} KB")
    print("\n  Section headings:")
    for h in stats["headings"]:
        print(f"    - {h}")

    print("\nAll structural assertions passed. Verdict: NO-DEPLOY.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
